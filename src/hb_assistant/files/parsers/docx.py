"""DOCXParser: bounded text/tables extraction (python-docx). No macro execution."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict

try:
    from docx import Document  # type: ignore
except ImportError:
    Document = None  # type: ignore


class DOCXParser:
    def parse(self, path: Path, max_chars: int = 8000) -> Dict[str, Any]:
        if Document is None:
            raise RuntimeError("python-docx not installed")
        try:
            doc = Document(str(path))
            parts: list[str] = []
            total = 0
            # paragraphs
            for para in doc.paragraphs:
                t = para.text or ""
                if t.strip():
                    parts.append(t)
                    total += len(t)
                    if total > max_chars:
                        break
            # tables (values only)
            if total < max_chars:
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join((cell.text or "").strip() for cell in row.cells)
                        if row_text.strip():
                            parts.append(row_text)
                            total += len(row_text)
                            if total > max_chars:
                                break
                    if total > max_chars:
                        break
            excerpt = "\n".join(parts)[:max_chars]
            return {
                "text_excerpt": excerpt,
                "char_count": len(excerpt),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
        except Exception as e:  # broad for encrypted/malformed etc
            msg = str(e)[:200]
            fc = "encrypted_or_password_protected" if "password" in msg.lower() or "encrypt" in msg.lower() else "parser_error"
            return {"text_excerpt": "", "char_count": 0, "error": msg, "failure_code": fc}
