"""Phase 10 — review-safe document/file parse read-model (local-only, raw-free).

Proves the read-model runs local parsers on synthetic fixtures and emits ONLY safe metadata (never the
extracted text), supports txt/md/docx/xlsx/pdf, degrades honestly on unsupported/missing files, and
that the `files parse-index` CLI verb works.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from typer.testing import CliRunner

from hb_assistant.cli.main import app
from hb_assistant.construction.second_brain.local_ai.file_parse_read_model import (
    build_file_index_read_model,
    build_file_parse_read_model,
    render_file_index_markdown,
)

runner = CliRunner()


def _make_fixtures(d: Path) -> dict[str, Path]:
    out: dict[str, Path] = {}
    (d / "note.txt").write_text("Synthetic text fixture.\nSecond line.\n", encoding="utf-8")
    out["txt"] = d / "note.txt"
    (d / "doc.md").write_text("# Synthetic\n\nBody paragraph.\n", encoding="utf-8")
    out["md"] = d / "doc.md"

    from docx import Document

    doc = Document()
    doc.add_paragraph("Synthetic docx paragraph.")
    doc.add_table(rows=1, cols=2)
    doc.save(str(d / "report.docx"))
    out["docx"] = d / "report.docx"

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws["A1"] = "synthetic"
    ws["B1"] = 123
    wb.save(str(d / "sheet.xlsx"))
    out["xlsx"] = d / "sheet.xlsx"

    from pypdf import PdfWriter

    w = PdfWriter()
    w.add_blank_page(width=200, height=200)
    with (d / "blank.pdf").open("wb") as fh:
        w.write(fh)
    out["pdf"] = d / "blank.pdf"
    return out


def test_read_model_is_raw_free_and_supported(tmp_path: Path) -> None:
    fx = _make_fixtures(tmp_path)
    for kind, p in fx.items():
        rm = build_file_parse_read_model(p)
        assert "text_excerpt" not in rm  # never the extracted text
        assert rm["parsed_status"] in ("parsed", "degraded"), (kind, rm)
        assert rm["extraction_method"] is not None
        assert rm["redaction"]["raw_text_excerpt_excluded"] is True
        assert rm["file_name"] == p.name
        # The hash covers the BOUNDED excerpt, not full text — named text_excerpt_hash + hash_scope.
        assert rm["hash_scope"] == "text_excerpt"
        assert "text_hash" not in rm  # legacy ambiguous field is gone (no alias)
        assert rm["text_excerpt_hash"] is None or rm["text_excerpt_hash"].startswith("sha256:")
    # txt fixture has content → an excerpt hash + positive length.
    txt_rm = build_file_parse_read_model(fx["txt"])
    assert txt_rm["text_length"] > 0
    assert txt_rm["text_excerpt_hash"].startswith("sha256:")
    assert txt_rm["hash_scope"] == "text_excerpt"


def test_unsupported_and_missing(tmp_path: Path) -> None:
    bad = tmp_path / "weird.xyz"
    bad.write_text("x", encoding="utf-8")
    rm = build_file_parse_read_model(bad)
    assert rm["parsed_status"] == "unsupported"
    assert rm["degraded_reason"].startswith("unsupported_extension")

    missing = build_file_parse_read_model(tmp_path / "nope.pdf")
    assert missing["parsed_status"] == "error"
    assert missing["degraded_reason"] == "file_not_found"


def test_index_and_markdown(tmp_path: Path) -> None:
    fx = _make_fixtures(tmp_path)
    index = build_file_index_read_model(list(fx.values()))
    assert index["counts"]["files"] == 5
    assert index["guardrails"]["no_raw_text_excerpt"] is True
    md = render_file_index_markdown(index)
    assert "File Parse Index" in md
    # No raw fixture content leaks into the index.
    blob = json.dumps(index) + md
    assert "Synthetic text fixture" not in blob
    assert "synthetic docx paragraph" not in blob.lower()


def test_cli_parse_index(tmp_path: Path) -> None:
    fx = _make_fixtures(tmp_path)
    res = runner.invoke(app, ["files", "parse-index", str(fx["txt"]), str(fx["xlsx"]), "--json"])
    assert res.exit_code == 0, res.output
    payload = json.loads(res.output)
    assert payload["counts"]["files"] == 2
    assert payload["guardrails"]["local_only"] is True


def test_cli_parse_index_no_json_emits_markdown(tmp_path: Path) -> None:
    # --no-json must be accepted (post-merge hardening) and print operator Markdown, not JSON.
    fx = _make_fixtures(tmp_path)
    res = runner.invoke(app, ["files", "parse-index", str(fx["txt"]), "--no-json"])
    assert res.exit_code == 0, res.output
    assert res.output.lstrip().startswith("# File Parse Index")
    with pytest.raises(json.JSONDecodeError):
        json.loads(res.output)
